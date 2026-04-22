# -*- coding: utf-8 -*-
import io
from typing import Any, Dict, Optional

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TAIGA_ACCENT = RGBColor(0x1E, 0x4D, 0x35)
TAIGA_TEXT = RGBColor(0x11, 0x18, 0x27)
TAIGA_MUTED = RGBColor(0x5C, 0x5B, 0x56)
TAIGA_BORDER_HEX = "E0DFD8"
TAIGA_CARD_BG_HEX = "F8FAF9"


def _fmt_num_int(v) -> str:
    try:
        return f"{float(v):,.0f}".replace(",", " ")
    except Exception:
        return str(v)


def _fmt_num_2(v) -> str:
    try:
        return f"{float(v):,.2f}".replace(",", " ")
    except Exception:
        return str(v)


def _fmt_eur(v, decimals=0) -> str:
    if decimals == 0:
        return f"EUR {_fmt_num_int(v)}"
    return f"EUR {_fmt_num_2(v)}"


def _to_number(x) -> float:
    if x is None:
        return 0.0
    if isinstance(x, (int, float)):
        return float(x)
    s = str(x).strip().replace("\xa0", " ").replace(" ", "")
    if not s:
        return 0.0
    if "," in s and "." in s:
        s = s.replace(",", "")
    elif "," in s:
        s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        return 0.0


def _shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_border(cell, color_hex: str = TAIGA_BORDER_HEX, size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    else:
        for el in list(tc_borders):
            tc_borders.remove(el)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tc_borders.append(el)


def _style_table(table) -> None:
    header = table.rows[0].cells
    for cell in header:
        _shade_cell(cell, "1E4D35")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            _set_cell_border(cell)
            if i % 2 == 0 and i != 0:
                _shade_cell(cell, TAIGA_CARD_BG_HEX)


def _add_section_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = TAIGA_TEXT


def _add_kv_table(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    _add_section_title(doc, title)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Value"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)
    _style_table(table)
    doc.add_paragraph("")


def _find_col(existing_cols, candidates):
    lower = {str(c).lower(): c for c in existing_cols}
    for cand in candidates:
        if cand in lower:
            return lower[cand]
    for cand in candidates:
        for lc, orig in lower.items():
            if cand in lc:
                return orig
    return None


def generate_offer_doc(
    payload: Dict[str, Any],
    products_df,
    leasing_info: Dict[str, Any],
    trad_summary: Dict[str, Any],
    logo_path: Optional[str] = None,
) -> bytes:
    doc = Document()

    if logo_path:
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(logo_path, width=Inches(2.1))
            p_logo.paragraph_format.space_after = Pt(24)
        except Exception:
            pass

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Taiga Forma - Commercial Offer")
    r.font.name = "Georgia"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = TAIGA_TEXT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Commercial summary including Taiga Forma pricing, Taiga Cycle and leasing view.")
    s.font.name = "Arial"
    s.font.size = Pt(10.5)
    s.font.color.rgb = TAIGA_MUTED
    sub.paragraph_format.space_after = Pt(16)

    params = payload.get("params", {}) or {}
    taiga = payload.get("taiga_forma", {}) or {}
    cycle = payload.get("taiga_cycle", {}) or {}

    _add_kv_table(doc, "Project Context", [
        ("Customer", payload.get("customer_name", "") or "-"),
        ("Project", payload.get("project_name", "") or "-"),
        ("Date", payload.get("date_str", "") or "-"),
    ])

    buyback_year = leasing_info.get("buyback_year", cycle.get("cycle_year", 0))
    _add_kv_table(doc, "Offer Summary", [
        ("Taiga Forma list price", _fmt_eur(taiga.get("list_price", trad_summary.get("taiga_list_price", 0.0)), 0)),
        ("Selected units", _fmt_num_int(taiga.get("units", 0))),
        ("Effective project area", f'{float(params.get("area_m2", 0.0)):,.2f} m2'),
        ("Taiga Cycle year", _fmt_num_int(cycle.get("cycle_year", 0))),
        ("Base monthly payment", _fmt_eur(leasing_info.get("monthly_price_base", 0.0), 0)),
        ("Monthly payment with Taiga Cycle", _fmt_eur(leasing_info.get("monthly_price_with_buyback", 0.0), 0)),
        ("Contract term", f'{int(leasing_info.get("term_months", 0))} months'),
        ("Buyback year in leasing", _fmt_num_int(buyback_year)),
    ])

    cols = ["Product", "Qty", "Unit price (EUR)", "Line total (EUR)", "Discount (%)", "Offer line (EUR)"]
    table = doc.add_table(rows=1, cols=len(cols))
    for i, col in enumerate(cols):
        table.rows[0].cells[i].text = col

    subtotal = 0.0
    offer_total = 0.0
    added_rows = 0

    if products_df is not None and not getattr(products_df, "empty", True):
        dfp = products_df.copy()
        name_col = _find_col(dfp.columns, ["product", "name", "product_name", "item", "description", "code", "sku"])
        qty_col = _find_col(dfp.columns, ["qty", "quantity", "kpl", "pcs", "amount"])
        unit_col = _find_col(dfp.columns, ["unit_price", "unit price", "unit_price_eur", "price"])
        disc_col = _find_col(dfp.columns, ["discount_pct", "discount %", "discount"])

        if "product" not in dfp.columns:
            dfp["product"] = dfp[name_col] if name_col is not None else dfp.index.astype(str)
        if "qty" not in dfp.columns and qty_col is not None:
            dfp["qty"] = dfp[qty_col]
        if "unit_price" not in dfp.columns and unit_col is not None:
            dfp["unit_price"] = dfp[unit_col]
        if "discount_pct" not in dfp.columns:
            dfp["discount_pct"] = dfp[disc_col] if disc_col is not None else 0.0

        for _, row in dfp.iterrows():
            name = str(row.get("product", ""))
            qty = _to_number(row.get("qty", 1))
            unit = _to_number(row.get("unit_price", 0))
            disc = _to_number(row.get("discount_pct", 0))
            if disc > 1.0:
                disc = disc / 100.0
            if qty > 0:
                line_total = qty * unit
                offer_line = line_total * (1.0 - disc)
                subtotal += line_total
                offer_total += offer_line

                cells = table.add_row().cells
                cells[0].text = name
                cells[1].text = _fmt_num_int(qty)
                cells[2].text = _fmt_eur(unit, 2)
                cells[3].text = _fmt_eur(line_total, 2)
                cells[4].text = _fmt_num_2(disc * 100) + " %"
                cells[5].text = _fmt_eur(offer_line, 2)
                added_rows += 1

    if added_rows == 0:
        row = table.add_row().cells
        row[0].text = "No products selected."
        for idx in range(1, len(cols)):
            row[idx].text = ""

    _style_table(table)
    doc.add_paragraph("")

    _add_kv_table(doc, "Commercial Totals", [
        ("Subtotal (list)", _fmt_eur(subtotal, 2)),
        ("Total offer price", _fmt_eur(offer_total, 2)),
    ])

    _add_kv_table(doc, "Leasing and Taiga Cycle", [
        ("Base monthly payment", _fmt_eur(leasing_info.get("monthly_price_base", 0.0), 2)),
        ("Monthly payment with Taiga Cycle", _fmt_eur(leasing_info.get("monthly_price_with_buyback", 0.0), 2)),
        ("Contract term", f'{int(leasing_info.get("term_months", 0))} months'),
        ("Taiga Cycle assumption", f'Applied at year {int(buyback_year)}' if int(buyback_year or 0) > 0 else "No Taiga Cycle applied"),
    ])

    _add_kv_table(doc, "Lifecycle Comparison", [
        ("Taiga Forma total present value", _fmt_eur(trad_summary.get("taiga_pv", 0.0), 0)),
        ("Traditional Model total present value", _fmt_eur(trad_summary.get("trad_pv", 0.0), 0)),
        ("Delta total present value", _fmt_eur(trad_summary.get("delta_pv", 0.0), 0)),
        ("Taiga Forma average lifecycle cost", f'EUR {_fmt_num_2(trad_summary.get("taiga_cost_m2_mo", 0.0))} / m2 / month'),
        ("Traditional Model average lifecycle cost", f'EUR {_fmt_num_2(trad_summary.get("trad_cost_m2_mo", 0.0))} / m2 / month'),
    ])

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
