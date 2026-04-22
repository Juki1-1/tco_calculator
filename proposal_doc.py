# proposal_doc.py
import io
import numbers
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TAIGA_ACCENT = RGBColor(0x1E, 0x4D, 0x35)
TAIGA_TEXT = RGBColor(0x11, 0x18, 0x27)
TAIGA_MUTED = RGBColor(0x5C, 0x5B, 0x56)
TAIGA_BORDER_HEX = "E0DFD8"
TAIGA_CARD_BG_HEX = "F8FAF9"


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TAIGA_TEXT

    for lvl, size in [(1, 20), (2, 14), (3, 11)]:
        style_name = f"Heading {lvl}"
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = "Georgia" if lvl == 1 else "Arial"
            s.font.size = Pt(size)
            s.font.bold = True
            s.font.color.rgb = TAIGA_TEXT
            s.paragraph_format.space_before = Pt(6)
            s.paragraph_format.space_after = Pt(6)

    if "TaigaCardHeading" not in doc.styles:
        s = doc.styles.add_style("TaigaCardHeading", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = "Arial"
        s.font.size = Pt(11)
        s.font.bold = True
        s.font.color.rgb = TAIGA_TEXT
        s.paragraph_format.space_before = Pt(4)
        s.paragraph_format.space_after = Pt(3)


def _set_page_margins(doc: Document, left=2.0, right=2.0, top=2.0, bottom=2.0) -> None:
    section = doc.sections[0]
    section.left_margin = Inches(left / 2.54)
    section.right_margin = Inches(right / 2.54)
    section.top_margin = Inches(top / 2.54)
    section.bottom_margin = Inches(bottom / 2.54)


def _fmt_num_int(v) -> str:
    try:
        return f"{float(v):,.0f}".replace(",", " ")
    except Exception:
        return str(v)


def _fmt_num_2(v) -> str:
    try:
        return f"{float(v):,.2f}".replace(",", " ")
    except Exception:
        return str(v)


def _fmt_eur(v, decimals=0) -> str:
    if decimals == 0:
        return f"EUR {_fmt_num_int(v)}"
    return f"EUR {_fmt_num_2(v)}"


def _shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_border(cell, color_hex: str = TAIGA_BORDER_HEX, size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    else:
        for el in list(tc_borders):
            tc_borders.remove(el)
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tc_borders.append(el)


def _style_table(table) -> None:
    header = table.rows[0].cells
    for cell in header:
        _shade_cell(cell, "1E4D35")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            _set_cell_border(cell)
            if i % 2 == 0 and i != 0:
                _shade_cell(cell, TAIGA_CARD_BG_HEX)


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def _add_kv_table(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    if not rows:
        return
    _add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Item"
    table.rows[0].cells[1].text = "Value"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)
    _style_table(table)
    doc.add_paragraph("")


def _add_pivot_table(doc: Document, title: str, df) -> None:
    if df is None or getattr(df, "empty", True):
        return
    _add_heading(doc, title, level=2)
    n_cols = len(df.columns) + 1
    table = doc.add_table(rows=1, cols=n_cols)
    header = table.rows[0].cells
    header[0].text = "Cost item"
    for j, col in enumerate(df.columns, start=1):
        header[j].text = str(col)
    for idx, row in df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(idx).replace("_", " ").title()
        for j, val in enumerate(row, start=1):
            text = _fmt_eur(val, 0) if isinstance(val, numbers.Number) else str(val)
            cells[j].text = text
    _style_table(table)
    doc.add_paragraph("")


def generate_proposal_doc(
    payload,
    df_pivot_taiga=None,
    df_pivot_trad=None,
    df_pivot_delta=None,
    locale: str = "fi_FI",
    logo_path: Optional[str] = "logo.png",
    **kwargs,
) -> bytes:
    if isinstance(df_pivot_taiga, str) and df_pivot_trad is None and df_pivot_delta is None:
        locale = df_pivot_taiga
        df_pivot_taiga = None

    doc = Document()
    _configure_styles(doc)
    _set_page_margins(doc)

    if logo_path:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(logo_path, width=Inches(2.1))
            p.paragraph_format.space_after = Pt(28)
        except Exception:
            pass

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Taiga Concept - Lifecycle Summary")
    r.font.name = "Georgia"
    r.font.size = Pt(25)
    r.font.bold = True
    r.font.color.rgb = TAIGA_TEXT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("Executive lifecycle summary for Taiga Forma, Taiga Cycle and the Traditional Model baseline.")
    s.font.name = "Arial"
    s.font.size = Pt(10.5)
    s.font.color.rgb = TAIGA_MUTED
    sub.paragraph_format.space_after = Pt(16)

    params = payload.get("params", {}) or {}
    results = payload.get("results", {}) or {}
    taiga = payload.get("taiga_forma", {}) or {}
    cycle = payload.get("taiga_cycle", {}) or {}
    trad = payload.get("traditional_model", {}) or {}
    leasing = payload.get("leasing", {}) or {}

    _add_kv_table(doc, "Project Context", [
        ("Customer", payload.get("customer_name", "") or "-"),
        ("Project", payload.get("project_name", "") or "-"),
        ("Date", payload.get("date_str", "") or "-"),
    ])

    _add_kv_table(doc, "Project Basis", [
        ("Horizon", f'{int(params.get("years", 0))} years'),
        ("WACC", f'{float(params.get("wacc", 0.0)) * 100:,.2f} %'),
        ("Shared project area", f'{float(params.get("shared_area_m2", params.get("area_m2", 0.0))):,.2f} m2'),
        ("Effective project area", f'{float(params.get("area_m2", 0.0)):,.2f} m2'),
        ("Energy intensity", f'{float(params.get("kwh_m2yr", 0.0)):,.2f} kWh / m2 / year'),
        ("Electricity price", f'EUR {float(params.get("elec_price", 0.0)):,.2f} / kWh'),
    ])

    _add_kv_table(doc, "Executive Summary", [
        ("Taiga Forma total present value", _fmt_eur(results.get("TCO_TAIGA_PV", 0.0), 0)),
        ("Traditional Model total present value", _fmt_eur(results.get("TCO_TRAD_PV", 0.0), 0)),
        ("Delta total present value", _fmt_eur(results.get("DIFF_TRAD_TAIGA", 0.0), 0)),
        ("Taiga Forma average lifecycle cost", f'EUR {_fmt_num_2(results.get("TAIGA_COST_M2_MONTH", 0.0))} / m2 / month'),
        ("Traditional Model average lifecycle cost", f'EUR {_fmt_num_2(results.get("TRAD_COST_M2_MONTH", 0.0))} / m2 / month'),
        ("Delta average lifecycle cost", f'EUR {_fmt_num_2(results.get("DELTA_COST_M2_MONTH", 0.0))} / m2 / month'),
    ])

    _add_kv_table(doc, "Taiga Forma", [
        ("Taiga Forma list price", _fmt_eur(taiga.get("list_price", 0.0), 0)),
        ("Selected units", _fmt_num_int(taiga.get("units", 0))),
        ("Effective area", f'{float(taiga.get("effective_area_m2", 0.0)):,.2f} m2'),
        ("Occupancy rate", f'{float(taiga.get("occupancy_rate", 0.0)) * 100:,.1f} %'),
        ("Standby share", f'{float(taiga.get("standby_share", 0.0)) * 100:,.1f} %'),
        ("Commissioning total", _fmt_eur(taiga.get("commissioning_total", 0.0), 0)),
        ("Maintenance total", _fmt_eur(taiga.get("maintenance_total", 0.0), 0)),
        ("End-of-life cost", _fmt_eur(taiga.get("end_of_life_cost", 0.0), 0)),
    ])

    _add_kv_table(doc, "Taiga Cycle", [
        ("Taiga Cycle year", _fmt_num_int(cycle.get("cycle_year", 0))),
    ])

    _add_kv_table(doc, "Traditional Model", [
        ("Investment price", f'EUR {_fmt_num_2(trad.get("investment_per_m2", 0.0))} / m2'),
        ("Closed rooms", _fmt_num_int(trad.get("room_qty", 0))),
        ("Run fraction", f'{float(trad.get("run_fraction", 0.0)) * 100:,.1f} %'),
        ("End-of-life", f'{float(trad.get("end_of_life_pct", 0.0)) * 100:,.1f} % of investment'),
        ("Commissioning total", _fmt_eur(trad.get("commissioning_total", 0.0), 0)),
        ("Maintenance total", _fmt_eur(trad.get("maintenance_total", 0.0), 0)),
    ])

    _add_kv_table(doc, "Leasing", [
        ("Base monthly payment", _fmt_eur(leasing.get("base_monthly", 0.0), 0)),
        ("Monthly payment with Taiga Cycle", _fmt_eur(leasing.get("monthly_with_buyback", 0.0), 0)),
        ("Contract term", f'{int(leasing.get("term_months", 0))} months'),
        ("Taiga Cycle year used in leasing", _fmt_num_int(leasing.get("buyback_year", 0))),
    ])

    doc.add_page_break()
    _add_pivot_table(doc, "Taiga Forma yearly breakdown (present value)", df_pivot_taiga)
    doc.add_page_break()
    _add_pivot_table(doc, "Traditional Model yearly breakdown (present value)", df_pivot_trad)
    doc.add_page_break()
    _add_pivot_table(doc, "Delta (Taiga Forma minus Traditional Model)", df_pivot_delta)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()
